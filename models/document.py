import docx
import re
import numpy as np
import os

class Document():
    def __init__(self, filename):
        self.filename = filename
        self.paragraphs = None

    def getPlainText(self):
        doc = self.getDocumentObj()
        text = ''
        for paragraph in doc.paragraphs:
            if not (paragraph.text == '' or paragraph.text.isspace()):
                text += paragraph.text + '\n'
        return text

    def getParagraphs(self):
        if type(self.paragraphs) is np.ndarray:
            return self.paragraphs
        p = []
        n_paragraphs = []
        doc = self.getDocumentObj()
        for paragraph in doc.paragraphs:
            if not (paragraph.text == '' or paragraph.text.isspace()):
                # remove initial and end tabs (if any)
                text = re.sub(r'^\t+|^\s+|\t+$|\s+$', '', paragraph.text, flags=re.M)
                # remove middle tabs (if any)
                text = re.sub(r'\t+|\s+', ' ', text, flags=re.M)
                p.append(text)
        n_paragraphs.append(len(p))
        return np.array(p)

    def saveParagraphs(self):
        self.paragraphs = self.getParagraphs()

    def getCleanParagraphs(self):
        paragraph_transformed = []
        paragraphs = self.getParagraphs()
        for p in paragraphs:
            # Lower case text
            text = p.lower()
            # Remove punctuation
            text = re.sub(r'\W+', ' ', text, flags=re.M)
            paragraph_transformed.append(text)
        return np.array(paragraph_transformed)

    def getCleanText(self):
        text = self.getPlainText() or ''
        # Lower case text
        text = text.lower()
        # Remove punctuation
        text = re.sub(r'\W+', ' ', text, flags=re.M)    
        return text
    
    # Invoke saveParagraphs() first
    def getParagraph(self, ix):
        if type(self.paragraphs) is np.ndarray:
            return self.paragraphs[ix]
        return -1

    # Invove saveParagraphs() first
    def getParagraphsLength(self):
        if type(self.paragraphs) is np.ndarray:
            return len(self.paragraphs)
        return -1

    def getDocumentObj(self):
        return docx.Document(self.filename)

    def getFilename(self):
        return os.path.basename(self.filename)
    
